"""Document loaders for different file types."""

import os
from typing import List
from langchain_community.document_loaders import PyPDFLoader, TextLoader
from langchain_core.documents import Document
from docx import Document as DocxDocument
from app.utils.logger import logger


def load_docx(file_path: str) -> List[Document]:
    """Load a DOCX document."""
    doc = DocxDocument(file_path)

    # Extract text from all paragraphs
    full_text = []
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            full_text.append(paragraph.text)

    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    full_text.append(cell.text)

    content = "\n\n".join(full_text)

    # Create a single document with all content
    # DOCX files don't have native page concept, so we treat it as one document
    document = Document(
        page_content=content,
        metadata={
            "source": file_path,
            "file_type": "docx"
        }
    )

    return [document]


async def load_document(file_path: str) -> List[Document]:
    """Load a document based on file extension."""
    _, ext = os.path.splitext(file_path)
    ext = ext.lower()

    try:
        if ext == ".pdf":
            loader = PyPDFLoader(file_path)
            documents = loader.load()
        elif ext in [".txt", ".md"]:
            loader = TextLoader(file_path)
            documents = loader.load()
        elif ext == ".docx":
            documents = load_docx(file_path)
        else:
            raise ValueError(f"Unsupported file type: {ext}")

        logger.info(f"Loaded {len(documents)} pages from {file_path}")
        return documents

    except Exception as e:
        logger.error(f"Error loading document {file_path}: {str(e)}")
        raise
